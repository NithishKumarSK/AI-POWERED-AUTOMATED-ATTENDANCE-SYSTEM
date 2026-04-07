import sys
import docx
from pypdf import PdfReader

try:
    doc = docx.Document('C14.docx')
    with open('C14_ext.txt', 'w', encoding='utf-8') as f:
        for para in doc.paragraphs:
            if str(para.text).strip():
                f.write(para.text + '\n')
                
    doc_final = docx.Document('FINAL DOCX PDF.docx')
    with open('FINAL_ext.txt', 'w', encoding='utf-8') as f:
        for para in doc_final.paragraphs:
            if str(para.text).strip():
                f.write(para.text + '\n')
                
    reader = PdfReader('C17.pdf')
    with open('C17_ext.txt', 'w', encoding='utf-8') as f:
        for page in reader.pages:
            t = page.extract_text()
            if t:
                f.write(t + '\n')
    print("DONE EXTRACTION!")
except Exception as e:
    print(f"Error: {e}")
